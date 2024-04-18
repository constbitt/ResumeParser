import os
import tempfile

import comtypes.client
import docx
import pdfkit
import markdown2
import shutil
from docx import Document


class FileConverter:
    def __init__(self):
        comtypes.CoInitialize()

    def __del__(self):
        comtypes.CoUninitialize()

    def convert_to_pdf(self, input_file, output_file):
        _, file_extension = os.path.splitext(input_file)
        file_extension = file_extension.lower()
        if file_extension == '.pdf':
            shutil.copyfile(input_file, output_file)
        if file_extension == '.txt':
            self.txt_pdf_converter(input_file, output_file)
        elif file_extension == '.rtf':
            self.rtf_pdf_converter(input_file, output_file)
        elif file_extension == '.docx':
            self.docx_pdf_converter(input_file, output_file)
        elif file_extension == '.html' or file_extension == '.htm':
            self.html_pdf_converter(input_file, output_file)
        elif file_extension == '.md':
            self.md_pdf_converter(input_file, output_file)
        elif file_extension == '.doc':
            self.doc_pdf_converter(input_file, output_file)
        else:
            print("Unsupported file format")

    def txt_pdf_converter(self, input_file, output_file):
        temp_dir = tempfile.gettempdir()
        tmp_file = os.path.join(temp_dir, "tmp.docx")
        self.txt_docx_converter(input_file, tmp_file)
        self.docx_pdf_converter(tmp_file, output_file)
        os.remove(tmp_file)

    def txt_docx_converter(self, input_file, output_file):
        with open(input_file, 'r', encoding='utf-8') as file:
            lines = file.readlines()
        doc = Document()
        for line in lines:
            doc.add_paragraph(line.strip())
        doc.save(output_file)

    def rtf_pdf_converter(self, input_file, output_file):
        wdFormatPDF = 17
        input_file = os.path.abspath(input_file)
        output_file = os.path.abspath(output_file)
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(input_file)
        doc.SaveAs(output_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

    def docx_pdf_converter(self, input_file, output_file):
        pdf_format = 17
        input_file = os.path.abspath(input_file)
        output_file = os.path.abspath(output_file)
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = docx.Document(input_file)
        in_file = word.Documents.Open(input_file)
        in_file.SaveAs(output_file, FileFormat=pdf_format)
        in_file.Close()
        word.Quit()

    def doc_pdf_converter(self, input_file, output_file):
        wdFormatPDF = 17
        word = comtypes.client.CreateObject("Word.Application")
        input_file = os.path.abspath(input_file)
        output_file = os.path.abspath(output_file)
        doc = word.Documents.Open(input_file)
        doc.SaveAs(output_file, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

    def html_pdf_converter(self, input_file, output_file):
        pdfkit.from_file(input_file, output_file)

    def md_pdf_converter(self, input_file, output_file):
        with open(input_file, 'r', encoding='utf-8') as file:
            markdown_content = file.read()
        html_content = markdown2.markdown(markdown_content)
        pdfkit.from_string(html_content, output_file)

